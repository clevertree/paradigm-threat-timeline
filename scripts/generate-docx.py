#!/usr/bin/env python3
"""
Generate a Word document (.docx) from all content/*.md files -> export/timeline.docx

Uses python-docx for document generation and markdown-it-py for markdown parsing.
Images are embedded using Pillow for pre-processing.

Requirements:  pip install python-docx markdown-it-py Pillow
"""

import io
import json
import os
import re
import sys
from pathlib import Path

# ---------------------------------------------------------------------------
# Dependencies
# ---------------------------------------------------------------------------
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("ERROR: pip install python-docx")

try:
    from markdown_it import MarkdownIt
except ImportError:
    sys.exit("ERROR: pip install markdown-it-py")

try:
    from PIL import Image
except ImportError:
    sys.exit("ERROR: pip install Pillow")

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR  = Path(__file__).resolve().parent
REPO_DIR    = SCRIPT_DIR.parent
CONTENT_DIR = REPO_DIR / "content"
EXPORT_DIR  = REPO_DIR / "export"

# Chapter prefixes that belong in the appendix volume
_APPENDIX_PREFIXES = ("16.", "17.")

with open(REPO_DIR / "package.json") as _f:
    PKG_VERSION = json.load(_f)["version"]


def _partition_files(md_files: list[Path]):
    """Split sorted md files into (book_files, appendix_files)."""
    book, appendix = [], []
    for p in md_files:
        if p.name.startswith(_APPENDIX_PREFIXES):
            appendix.append(p)
        else:
            book.append(p)
    return book, appendix

# ---------------------------------------------------------------------------
# Roman numerals & tier classifier (shared with generate-pdf.py logic)
# ---------------------------------------------------------------------------
_ROMAN = [
    "", "I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X",
    "XI", "XII", "XIII", "XIV", "XV", "XVI", "XVII", "XVIII", "XIX", "XX",
]

_NUM_RE = re.compile(r"^(\d{2})\.(\d{2})\.(\d{2})")

def _classify_tier(filename: str):
    """Return (tier, xx, yy, zz) from an ``XX.YY.ZZ-slug.md`` filename."""
    m = _NUM_RE.match(filename)
    if not m:
        return ("part", "00", "00", "00")
    xx, yy, zz = m.group(1), m.group(2), m.group(3)
    if yy == "00" and zz == "00":
        return ("part", xx, yy, zz)
    if zz == "00":
        return ("chapter", xx, yy, zz)
    return ("subsection", xx, yy, zz)

# ---------------------------------------------------------------------------
# Page dimensions: 6 x 9 inch trade paperback
# ---------------------------------------------------------------------------
PAGE_W = Inches(6)
PAGE_H = Inches(9)
MARGIN_LR     = Cm(2.0)    # equal left/right margins (print services add gutter)
MARGIN_TOP     = Cm(1.8)
MARGIN_BOTTOM  = Cm(2.0)
BOOK_TITLE     = "PARADIGM THREAT: THE THIRD STORY"

# ---------------------------------------------------------------------------
# Image settings
# ---------------------------------------------------------------------------
MAX_IMG_PX    = 1200
JPEG_QUALITY  = 82
MAX_IMG_IN    = 4.0    # max image width in inches (fits 6x9 with margins)

# ---------------------------------------------------------------------------
# Token attribute helper (markdown-it-py can use dict or list-of-tuples)
# ---------------------------------------------------------------------------
def _attr(tok, key: str, default: str = "") -> str:
    if tok.attrs is None:
        return default
    if isinstance(tok.attrs, dict):
        return tok.attrs.get(key, default)
    for k, v in tok.attrs:
        if k == key:
            return v
    return default

# ---------------------------------------------------------------------------
# Plain text extractor from inline token tree
# ---------------------------------------------------------------------------
def _plain(tok) -> str:
    parts = []
    if tok.children:
        for c in tok.children:
            if c.type in ("text", "code_inline"):
                parts.append(c.content)
            elif c.type == "softbreak":
                parts.append(" ")
            elif c.type == "hardbreak":
                parts.append("\n")
            elif c.children:
                parts.append(_plain(c))
    elif tok.content:
        parts.append(tok.content)
    return "".join(parts)

# ---------------------------------------------------------------------------
# Table of Contents (manually built — mirrors PDF _render_toc output)
# ---------------------------------------------------------------------------
def _add_toc(doc: Document, md_files: list):
    """Build TOC entries manually from the content file list.

    This produces styled paragraphs (TOC 1/2/3) that match the PDF's
    rendered Table of Contents.  Because python-docx cannot know final
    page numbers, we omit them — the PDF is the page-numbered copy.
    """
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = heading.add_run("Table of Contents")
    hr.bold = True
    hr.font.size = Pt(22)
    hr.font.name = "Libre Franklin"
    hr.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    # Gold rule below heading
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rule.add_run("\u2500" * 60)
    rr.font.size = Pt(7)
    rr.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)
    rule.paragraph_format.space_after = Pt(8)

    chapter_counter = 0

    for path in md_files:
        tier, xx, yy, zz = _classify_tier(path.name)

        # Read H1 title
        try:
            with open(path, encoding="utf-8") as f:
                first_line = f.readline().strip()
            title = first_line.lstrip("#").strip() if first_line.startswith("#") else path.stem
        except Exception:
            title = path.stem

        if tier == "part":
            roman = _ROMAN[int(xx)] if int(xx) < len(_ROMAN) else str(xx)
            display = f"Part {roman}: {title}"

            # Split century-based part titles into main + subtitle
            pm = re.match(
                r"^(\d{1,2}(?:st|nd|rd|th)\s+Century(?:\s+C\.E\.)?)"
                r"(?:\s*[.:\-—]\s*|\s+)(.+)$",
                title,
            )
            if pm:
                main_title = f"Part {roman}: {pm.group(1).strip().rstrip('.')}"
                subtitle   = pm.group(2).strip().rstrip(".")
            else:
                main_title = display
                subtitle   = ""

            para = doc.add_paragraph(style="TOC 1")
            run = para.add_run(main_title)
            run.bold = True
            run.font.name = "Libre Franklin"
            run.font.size = Pt(10.5)

            if subtitle:
                sub_para = doc.add_paragraph(style="TOC 1")
                sub_para.paragraph_format.space_before = Pt(0)
                sub_para.paragraph_format.left_indent = Cm(0.4)
                sr = sub_para.add_run(subtitle)
                sr.italic = True
                sr.bold = False
                sr.font.name = "Libre Franklin"
                sr.font.size = Pt(8.8)
                sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        elif tier == "chapter":
            chapter_counter += 1
            para = doc.add_paragraph(style="TOC 2")
            cr = para.add_run(f"Chapter {chapter_counter}: {title}")
            cr.font.name = "EB Garamond"
            cr.font.size = Pt(9.5)

        elif tier == "subsection":
            para = doc.add_paragraph(style="TOC 3")
            sr = para.add_run(title)
            sr.italic = True
            sr.font.name = "EB Garamond"
            sr.font.size = Pt(8.5)
            sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

# ---------------------------------------------------------------------------
# DOCX Renderer
# ---------------------------------------------------------------------------
class DOCXRenderer:
    def __init__(self):
        self._img_cache: dict = {}
        self._reset_doc()

    def _reset_doc(self):
        """Create a fresh Document with styles and layout."""
        self.doc = Document()
        self.md  = MarkdownIt("commonmark", {"typographer": False}).enable("table")
        self._chapter_counter = 0
        self._current_part_xx = None
        self._current_part_label = ""   # e.g. "Part I — Introduction"
        self._current_chapter_title = ""  # e.g. "The Third Story"
        self._setup_styles()
        self._setup_page_layout()

    # Font families — match the PDF's EB Garamond + Libre Franklin + JetBrains Mono
    FONT_SERIF = "EB Garamond"
    FONT_SANS  = "Libre Franklin"
    FONT_MONO  = "JetBrains Mono"

    # Secondary text size for lists, tables, blockquotes, code
    SECONDARY_SIZE = Pt(9.5)

    def _setup_styles(self):
        doc = self.doc

        # Body text default — EB Garamond 11pt (matches PDF)
        style = doc.styles["Normal"]
        style.font.name = self.FONT_SERIF
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Enable auto-hyphenation so justified text doesn't get huge word gaps
        pPr = style.element.get_or_add_pPr()
        hyph = OxmlElement("w:suppressAutoHyphens")
        hyph.set(qn("w:val"), "0")          # 0 = allow hyphenation
        pPr.append(hyph)
        # Set East-Asian / complex-script fallback to same font
        rpr = style.element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rpr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), self.FONT_SERIF)
        rFonts.set(qn("w:cs"), self.FONT_SERIF)
        # Set explicit language so LibreOffice applies hyphenation dictionary
        lang = rpr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            rpr.append(lang)
        lang.set(qn("w:val"), "en-US")
        lang.set(qn("w:eastAsia"), "en-US")
        lang.set(qn("w:bidi"), "en-US")

        # Heading styles — Libre Franklin (sans) for H1-H3, sizes matching PDF
        h_specs = {
            "Heading 1": (20, True, RGBColor(0x1a, 0x1a, 0x1a), self.FONT_SANS),
            "Heading 2": (14, True, RGBColor(0x2c, 0x2c, 0x2c), self.FONT_SANS),
            "Heading 3": (11, True, RGBColor(0x3a, 0x3a, 0x3a), self.FONT_SANS),
            "Heading 4": (10.5, True, RGBColor(0x4e, 0x4e, 0x4e), self.FONT_SERIF),
        }
        for sname, (sz, bold, clr, fname) in h_specs.items():
            s = doc.styles[sname]
            s.font.name  = fname
            s.font.size  = Pt(sz)
            s.font.bold  = bold
            s.font.color.rgb = clr
            s.paragraph_format.space_before = Pt(14 if sz >= 16 else 10)
            s.paragraph_format.space_after  = Pt(4)

        # Quote style
        if "Quote" not in [s.name for s in doc.styles]:
            qs = doc.styles.add_style("Quote", WD_STYLE_TYPE.PARAGRAPH)
        else:
            qs = doc.styles["Quote"]
        qs.base_style = doc.styles["Normal"]
        qs.font.name   = self.FONT_SERIF
        qs.font.italic = True
        qs.font.size   = self.SECONDARY_SIZE
        qs.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        qs.paragraph_format.left_indent  = Cm(1.2)
        qs.paragraph_format.space_before = Pt(4)
        qs.paragraph_format.space_after  = Pt(4)

        # List styles — smaller font, tighter spacing
        for ls_name in ("List Bullet", "List Number"):
            try:
                ls = doc.styles[ls_name]
            except KeyError:
                continue
            ls.font.name = self.FONT_SERIF
            ls.font.size = self.SECONDARY_SIZE
            ls.paragraph_format.space_before = Pt(1)
            ls.paragraph_format.space_after  = Pt(1)

        # Caption style
        if "Caption" not in [s.name for s in doc.styles]:
            cs = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        else:
            cs = doc.styles["Caption"]
        cs.base_style = doc.styles["Normal"]
        cs.font.name   = self.FONT_SANS
        cs.font.italic = True
        cs.font.size   = Pt(9)
        cs.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        cs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cs.paragraph_format.space_before = Pt(2)
        cs.paragraph_format.space_after  = Pt(8)

        # TOC styles — match PDF hierarchy (Sans bold for Parts, Serif for rest)
        # Parts: Libre Franklin Bold 10.5pt, dark, no indent, extra spacing
        for toc_name in ("TOC 1", "TOC 2", "TOC 3"):
            try:
                _ = doc.styles[toc_name]
            except KeyError:
                doc.styles.add_style(toc_name, WD_STYLE_TYPE.PARAGRAPH)

        toc1 = doc.styles["TOC 1"]
        toc1.base_style = doc.styles["Normal"]
        toc1.font.name = self.FONT_SANS
        toc1.font.size = Pt(10.5)
        toc1.font.bold = True
        toc1.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        toc1.paragraph_format.space_before = Pt(6)
        toc1.paragraph_format.space_after  = Pt(2)
        toc1.paragraph_format.left_indent  = Cm(0)
        toc1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Chapters: EB Garamond 9.5pt, brownish, indent ~8mm
        toc2 = doc.styles["TOC 2"]
        toc2.base_style = doc.styles["Normal"]
        toc2.font.name = self.FONT_SERIF
        toc2.font.size = Pt(9.5)
        toc2.font.bold = False
        toc2.font.color.rgb = RGBColor(0x55, 0x3a, 0x1a)
        toc2.paragraph_format.space_before = Pt(1)
        toc2.paragraph_format.space_after  = Pt(1)
        toc2.paragraph_format.left_indent  = Cm(0.8)
        toc2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Subsections: EB Garamond Italic 8.5pt, gray, indent ~14mm
        toc3 = doc.styles["TOC 3"]
        toc3.base_style = doc.styles["Normal"]
        toc3.font.name = self.FONT_SERIF
        toc3.font.size = Pt(8.5)
        toc3.font.bold = False
        toc3.font.italic = True
        toc3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        toc3.paragraph_format.space_before = Pt(0)
        toc3.paragraph_format.space_after  = Pt(0)
        toc3.paragraph_format.left_indent  = Cm(1.4)
        toc3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ---------------------------------------------------------------
    #  Page layout: 6x9, equal margins, running headers
    # ---------------------------------------------------------------
    def _setup_page_layout(self):
        """Configure the initial section for 6×9 trade paperback with
        equal margins and running headers/footers."""
        section = self.doc.sections[0]
        self._apply_section_dims(section)

        settings = self.doc.settings.element

        # Enable auto-hyphenation (document-level) for better justified text
        auto_hyp = OxmlElement("w:autoHyphenation")
        auto_hyp.set(qn("w:val"), "1")
        settings.append(auto_hyp)

        # Set document-level default language for hyphenation
        styles_el = self.doc.styles.element
        doc_defaults = styles_el.find(qn("w:docDefaults"))
        if doc_defaults is None:
            doc_defaults = OxmlElement("w:docDefaults")
            styles_el.insert(0, doc_defaults)
        rpr_default = doc_defaults.find(qn("w:rPrDefault"))
        if rpr_default is None:
            rpr_default = OxmlElement("w:rPrDefault")
            doc_defaults.append(rpr_default)
        rpr_el = rpr_default.find(qn("w:rPr"))
        if rpr_el is None:
            rpr_el = OxmlElement("w:rPr")
            rpr_default.append(rpr_el)
        def_lang = OxmlElement("w:lang")
        def_lang.set(qn("w:val"), "en-US")
        def_lang.set(qn("w:eastAsia"), "en-US")
        def_lang.set(qn("w:bidi"), "en-US")
        rpr_el.append(def_lang)

        # Title page: suppress headers/footers
        section.different_first_page_header_footer = True

    def _apply_section_dims(self, section):
        """Set page size and margins on a section."""
        section.page_width  = PAGE_W
        section.page_height = PAGE_H
        section.top_margin    = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin   = MARGIN_LR
        section.right_margin  = MARGIN_LR

    @staticmethod
    def _add_page_number_field(para):
        """Insert a PAGE field code into a paragraph."""
        run = para.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)

        run2 = para.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run2._r.append(instr)

        run3 = para.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run3._r.append(fld_end)

    def _set_header_footer(self, section, chapter_title: str):
        """Set up running headers and page-number footers.

        Header:  page#  ........  CHAPTER TITLE
        Footer:  centred page number
        """
        # Calculate usable width for tab stop
        usable = PAGE_W - MARGIN_LR - MARGIN_LR
        tab_right = usable  # right-aligned tab at full usable width

        # ── Header ──
        hdr = section.header
        hdr.is_linked_to_previous = False
        for p in hdr.paragraphs:
            p.clear()
        p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        pf = p.paragraph_format
        pf.tab_stops.add_tab_stop(tab_right, WD_TAB_ALIGNMENT.RIGHT)
        pf.space_after = Pt(0)
        # Page number left
        self._add_page_number_field(p)
        for r in p.runs:
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        # Tab + chapter title right
        r_title = p.add_run(f"\t{chapter_title.upper()}")
        r_title.font.size = Pt(8.5)
        r_title.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # ── Footer: centred page number ──
        ftr = section.footer
        ftr.is_linked_to_previous = False
        for p in ftr.paragraphs:
            p.clear()
        p = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        self._add_page_number_field(p)
        for r in p.runs:
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    def _new_section_with_headers(self, chapter_title: str):
        """Add a new section (page break) with updated running headers."""
        # Add section break (new page)
        new_sec = self.doc.add_section()
        self._apply_section_dims(new_sec)
        new_sec.different_first_page_header_footer = True
        self._set_header_footer(new_sec, chapter_title)
        self._current_chapter_title = chapter_title

    # ---------------------------------------------------------------
    #  Image handling
    # ---------------------------------------------------------------
    def _prepare_image(self, src: str):
        if src in self._img_cache:
            return self._img_cache[src]

        if src.startswith(("http://", "https://")):
            self._img_cache[src] = None
            return None

        rel = src.lstrip("/")
        abs_path = REPO_DIR / rel
        if not abs_path.is_file():
            print(f"    WARN image not found: {rel}", file=sys.stderr)
            self._img_cache[src] = None
            return None

        try:
            img = Image.open(abs_path)
            if img.width > MAX_IMG_PX:
                r = MAX_IMG_PX / img.width
                img = img.resize((MAX_IMG_PX, int(img.height * r)), Image.LANCZOS)

            if img.mode in ("RGBA", "P", "LA"):
                bg = Image.new("RGB", img.size, (255, 255, 255))
                if img.mode == "P":
                    img = img.convert("RGBA")
                bg.paste(img, mask=img.split()[-1] if "A" in img.mode else None)
                img = bg
            elif img.mode != "RGB":
                img = img.convert("RGB")

            buf = io.BytesIO()
            img.save(buf, "JPEG", quality=JPEG_QUALITY, optimize=True)
            buf.seek(0)

            # Compute display width in inches, respecting max
            orig_w_in = img.width / 96   # assume 96 dpi
            display_w = min(orig_w_in, MAX_IMG_IN)

            self._img_cache[src] = (buf, display_w)
            return self._img_cache[src]
        except Exception as exc:
            print(f"    WARN image failed ({rel}): {exc}", file=sys.stderr)
            self._img_cache[src] = None
            return None

    def _embed_image(self, src: str, caption: str = ""):
        result = self._prepare_image(src)
        if not result:
            return
        buf, width_in = result
        buf.seek(0)

        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(buf, width=Inches(width_in))

        if caption:
            cp = self.doc.add_paragraph(style="Caption")
            cp.add_run(caption)

    # ---------------------------------------------------------------
    #  Inline run builder (bold / italic / code / link)
    # ---------------------------------------------------------------
    def _apply_inline(self, para, children, font_size_override=None):
        if not children:
            return
        bold = False
        italic = False
        link_text_parts: list = []
        link_href: str = ""
        in_link = False

        def _style_run(run):
            """Apply font_size_override to a run if set."""
            if font_size_override:
                run.font.size = font_size_override

        for tok in children:
            tp = tok.type

            if tp == "text":
                if in_link:
                    link_text_parts.append(tok.content)
                else:
                    run = para.add_run(tok.content)
                    run.bold   = bold
                    run.italic = italic
                    _style_run(run)

            elif tp == "softbreak":
                if in_link:
                    link_text_parts.append(" ")
                else:
                    run = para.add_run(" ")
                    run.bold = bold; run.italic = italic
                    _style_run(run)

            elif tp == "hardbreak":
                run = para.add_run("\n")
                _style_run(run)

            elif tp == "code_inline":
                run = para.add_run(tok.content)
                run.font.name = self.FONT_MONO
                run.font.size = font_size_override or Pt(10)

            elif tp == "strong_open":
                bold = True
            elif tp == "strong_close":
                bold = False
            elif tp == "em_open":
                italic = True
            elif tp == "em_close":
                italic = False

            elif tp == "link_open":
                link_href = _attr(tok, "href")
                link_text_parts = []
                in_link = True
            elif tp == "link_close":
                link_text = "".join(link_text_parts)
                run = para.add_run(link_text)
                run.bold   = bold
                run.italic = italic
                run.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)  # gold
                _style_run(run)
                try:
                    self._add_hyperlink(para, run, link_href)
                except Exception:
                    pass
                in_link = False
                link_text_parts = []

            elif tp == "image":
                isrc = _attr(tok, "src")
                ialt = _attr(tok, "alt") or tok.content or ""
                self._embed_image(isrc, ialt)

            elif tp == "html_inline":
                if "<br" in (tok.content or ""):
                    para.add_run("\n")

            else:
                if tok.children:
                    self._apply_inline(para, tok.children)
                elif tok.content:
                    run = para.add_run(tok.content)
                    run.bold = bold; run.italic = italic

    def _add_hyperlink(self, para, run, url: str):
        """Replace the last run with a proper Word hyperlink field."""
        # python-docx doesn't support adding hyperlinks to existing runs easily;
        # we add the relationship and mark the run with it
        part = para.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        hyperlink.append(run._r)
        para._p.append(hyperlink)

    # ---------------------------------------------------------------
    #  Block rendering
    # ---------------------------------------------------------------
    def _render_heading(self, level: int, inline_tok):
        text  = _plain(inline_tok)
        style = f"Heading {min(level, 4)}"
        para  = self.doc.add_paragraph(style=style)
        para.add_run(text)

    def _render_paragraph(self, inline_tok):
        if not inline_tok.children:
            if inline_tok.content:
                self.doc.add_paragraph(inline_tok.content)
            return

        # Pure image paragraph
        real = [c for c in inline_tok.children
                if c.type not in ("softbreak", "hardbreak")]
        if len(real) == 1 and real[0].type == "image":
            img = real[0]
            self._embed_image(_attr(img, "src"), _attr(img, "alt") or img.content or "")
            return

        # Image + caption pattern
        if (len(real) >= 2 and real[0].type == "image" and real[1].type == "em_open"):
            img = real[0]
            cap_parts = []
            for c in real[2:]:
                if c.type == "em_close": break
                if c.type == "text": cap_parts.append(c.content)
            caption = "".join(cap_parts) or _attr(img, "alt") or img.content or ""
            self._embed_image(_attr(img, "src"), caption)
            return

        para = self.doc.add_paragraph()
        self._apply_inline(para, inline_tok.children)

    def _render_blockquote(self, inner_tokens: list):
        # Collect text runs from nested tokens into quote paragraphs
        for tok in inner_tokens:
            if tok.type == "inline":
                para = self.doc.add_paragraph(style="Quote")
                self._apply_inline(para, tok.children,
                                   font_size_override=self.SECONDARY_SIZE)
            elif tok.type in ("paragraph_open", "paragraph_close",
                              "blockquote_open", "blockquote_close"):
                pass

    def _render_list(self, items: list, ordered: bool, depth: int = 0):
        word_style = "List Number" if ordered else "List Bullet"
        for item_tokens in items:
            # Flatten inline content from the list item
            for tok in item_tokens:
                if tok.type == "inline":
                    para = self.doc.add_paragraph(style=word_style)
                    para.paragraph_format.space_before = Pt(1)
                    para.paragraph_format.space_after  = Pt(1)
                    if depth > 0:
                        para.paragraph_format.left_indent = Cm(depth * 1.0)
                    self._apply_inline(para, tok.children,
                                       font_size_override=self.SECONDARY_SIZE)
                elif tok.type == "bullet_list_open":
                    # nested list
                    pass

    def _render_table(self, header_row: list, body_rows: list):
        cols = len(header_row)
        if cols == 0:
            return

        TABLE_FONT_SIZE = Pt(9)

        table = self.doc.add_table(rows=1, cols=cols)
        table.style = "Table Grid"
        table.autofit = False

        # ── Compute proportional column widths based on content length ──
        # Gather all text per column (header + body) to estimate widths
        col_text_len = [0] * cols
        for j, cell_tokens in enumerate(header_row):
            for tok in cell_tokens:
                if tok.type == "inline":
                    for child in (tok.children or []):
                        col_text_len[j] += len(child.content or "")
        for row_cells in body_rows:
            for j, cell_tokens in enumerate(row_cells[:cols]):
                for tok in cell_tokens:
                    if tok.type == "inline":
                        for child in (tok.children or []):
                            col_text_len[j] += len(child.content or "")

        # Usable width = page width minus margins
        usable = PAGE_W - MARGIN_LR - MARGIN_LR
        # Minimum column width so nothing collapses
        min_col = Inches(0.6)
        total_text = max(sum(col_text_len), 1)
        col_widths = []
        for j in range(cols):
            proportion = col_text_len[j] / total_text
            w = int(usable * proportion)
            if w < min_col:
                w = min_col
            col_widths.append(w)
        # Normalize so widths sum to usable
        w_sum = sum(col_widths)
        if w_sum > 0:
            col_widths = [int(w * usable / w_sum) for w in col_widths]

        # Apply column widths to all rows
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                cell.width = col_widths[j]
        # Also set on table column objects
        for j, col in enumerate(table.columns):
            col.width = col_widths[j]

        # ── Compact cell margins via table-level cellMargin ──
        tbl_pr = table._tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tbl_pr)
        cm_el = OxmlElement("w:tblCellMar")
        for side, val in [("w:top", "30"), ("w:bottom", "30"),
                          ("w:left", "60"), ("w:right", "60")]:
            el = OxmlElement(side)
            el.set(qn("w:w"), val)
            el.set(qn("w:type"), "dxa")
            cm_el.append(el)
        tbl_pr.append(cm_el)

        # ── Header row ──
        hdr_cells = table.rows[0].cells
        for j, cell_tokens in enumerate(header_row):
            p = hdr_cells[j].paragraphs[0]
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            for tok in cell_tokens:
                if tok.type == "inline":
                    for child in (tok.children or []):
                        run = p.add_run(child.content if child.content else "")
                        run.bold = True
                        run.font.size = TABLE_FONT_SIZE
                        run.font.name = self.FONT_SANS
            # Header shading
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F0F0F0")
            hdr_cells[j]._tc.get_or_add_tcPr().append(shd)

        # ── Body rows ──
        for row_cells in body_rows:
            row = table.add_row().cells
            for j, cell_tokens in enumerate(row_cells[:cols]):
                p = row[j].paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                row[j].width = col_widths[j]
                for tok in cell_tokens:
                    if tok.type == "inline":
                        self._apply_inline(p, tok.children or [],
                                           font_size_override=TABLE_FONT_SIZE)

        self.doc.add_paragraph()  # spacing after table

    def _render_code_block(self, content: str):
        # Use a shaded paragraph with monospace font
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(1)
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after  = Pt(3)
        run = para.add_run(content.rstrip())
        run.font.name = self.FONT_MONO
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        # Light grey shading
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F5F5F5")
        pPr.append(shd)

    # ---------------------------------------------------------------
    #  Token stream walker
    # ---------------------------------------------------------------
    def _render_tokens(self, tokens):
        i = 0
        while i < len(tokens):
            tok = tokens[i]
            tp  = tok.type

            if tp == "heading_open":
                level = int(tok.tag[1]) if tok.tag else 2
                inline_tok = tokens[i + 1] if i + 1 < len(tokens) else None
                if inline_tok and inline_tok.type == "inline":
                    self._render_heading(level, inline_tok)
                    i += 2  # skip heading_close
                else:
                    i += 1

            elif tp == "paragraph_open":
                inline_tok = tokens[i + 1] if i + 1 < len(tokens) else None
                if inline_tok and inline_tok.type == "inline":
                    self._render_paragraph(inline_tok)
                    i += 2  # skip paragraph_close
                else:
                    i += 1

            elif tp == "blockquote_open":
                # collect until matching blockquote_close
                depth = 1
                j = i + 1
                while j < len(tokens) and depth > 0:
                    if tokens[j].type == "blockquote_open":  depth += 1
                    if tokens[j].type == "blockquote_close": depth -= 1
                    j += 1
                self._render_blockquote(tokens[i + 1: j - 1])
                i = j

            elif tp in ("bullet_list_open", "ordered_list_open"):
                ordered = tp == "ordered_list_open"
                depth_counter = 1
                j = i + 1
                items: list = []
                current_item: list = []
                while j < len(tokens):
                    jt = tokens[j].type
                    if jt in ("bullet_list_open", "ordered_list_open"):
                        depth_counter += 1
                    elif jt in ("bullet_list_close", "ordered_list_close"):
                        depth_counter -= 1
                        if depth_counter == 0:
                            if current_item:
                                items.append(current_item)
                            j += 1
                            break
                    elif jt == "list_item_open":
                        current_item = []
                    elif jt == "list_item_close":
                        items.append(current_item)
                        current_item = []
                    else:
                        current_item.append(tokens[j])
                    j += 1
                self._render_list(items, ordered)
                i = j

            elif tp == "fence" or tp == "code_block":
                self._render_code_block(tok.content)
                i += 1

            elif tp == "hr":
                para = self.doc.add_paragraph()
                pPr = para._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"),   "single")
                bottom.set(qn("w:sz"),    "6")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "CCCCCC")
                pBdr.append(bottom)
                pPr.append(pBdr)
                i += 1

            elif tp == "table_open":
                # Collect header and body rows
                header_row: list = []
                body_rows:  list = []
                in_head = False
                in_body = False
                current_row: list = []
                current_cell: list = []
                j = i + 1
                while j < len(tokens) and tokens[j].type != "table_close":
                    tt = tokens[j].type
                    if tt == "thead_open":   in_head = True
                    elif tt == "thead_close": in_head = False
                    elif tt == "tbody_open":  in_body = True
                    elif tt == "tbody_close": in_body = False
                    elif tt == "tr_open":     current_row = []
                    elif tt == "tr_close":
                        if in_head: header_row = current_row
                        else:       body_rows.append(current_row)
                    elif tt in ("th_open", "td_open"): current_cell = []
                    elif tt in ("th_close", "td_close"):
                        current_row.append(current_cell)
                    else:
                        current_cell.append(tokens[j])
                    j += 1
                self._render_table(header_row, body_rows)
                i = j + 1  # skip table_close

            elif tp == "html_block":
                # Skip raw HTML blocks (iframes etc.)
                i += 1

            else:
                i += 1

    # ---------------------------------------------------------------
    #  Strip first H1 from a token stream
    # ---------------------------------------------------------------
    @staticmethod
    def _strip_first_h1(tokens: list) -> list:
        """Remove the first H1 heading (open + inline + close) from tokens."""
        out = []
        skipped = False
        i = 0
        while i < len(tokens):
            if (not skipped and tokens[i].type == "heading_open"
                    and tokens[i].tag == "h1"):
                skipped = True
                i += 1
                while i < len(tokens) and tokens[i].type != "heading_close":
                    i += 1
                i += 1
                continue
            out.append(tokens[i])
            i += 1
        return out

    # ---------------------------------------------------------------
    #  Heading-demoted token renderer (for sub-sections)
    # ---------------------------------------------------------------
    def _render_tokens_demoted(self, tokens, demote: int = 1):
        """Like _render_tokens but demotes heading levels by *demote*."""
        for tok in tokens:
            if tok.type == "heading_open" and tok.tag and tok.tag[0] == "h":
                old_level = int(tok.tag[1])
                new_level = min(old_level + demote, 6)
                tok.tag = f"h{new_level}"
        self._render_tokens(tokens)

    # ---------------------------------------------------------------
    #  Article entry point  (3-tier: Part / Chapter / Sub-section)
    # ---------------------------------------------------------------
    def render_article(self, md_path: Path):
        text   = md_path.read_text(encoding="utf-8")
        tokens = self.md.parse(text)

        tier, xx, yy, zz = _classify_tier(md_path.name)
        part_num = int(xx)
        roman = _ROMAN[part_num] if part_num < len(_ROMAN) else str(part_num)

        # Extract H1 for labelling
        h1_match = re.match(r"^#\s+(.+)", text, re.MULTILINE)
        h1_title = h1_match.group(1).strip() if h1_match else md_path.stem

        if tier == "part":
            # ── PART ──────────────────────────────────────────────
            self._current_part_xx = xx
            self._current_part_label = f"Part {roman} — {h1_title}"

            # New section with updated running header
            self._new_section_with_headers(h1_title)

            # Part title heading
            h = self.doc.add_paragraph(style="Heading 1")
            run = h.add_run(f"Part {roman} — {h1_title}")
            run.bold = True

            tokens = self._strip_first_h1(tokens)
            self._render_tokens(tokens)
            self.doc.add_paragraph()

        elif tier == "chapter":
            # ── CHAPTER ───────────────────────────────────────────
            self._chapter_counter += 1
            ch_num = self._chapter_counter

            # New section with updated running header
            self._new_section_with_headers(h1_title)

            # Chapter heading
            h = self.doc.add_paragraph(style="Heading 2")
            run = h.add_run(f"Chapter {ch_num}: {h1_title}")
            run.bold = True

            tokens = self._strip_first_h1(tokens)
            self._render_tokens(tokens)
            self.doc.add_paragraph()

        else:
            # ── SUB-SECTION ───────────────────────────────────────
            # Horizontal rule separator (not a page break)
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = para.add_run("— ✦ —")
            r.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)
            r.font.size = Pt(12)

            tokens = self._strip_first_h1(tokens)
            self._render_tokens_demoted(tokens, demote=1)
            self.doc.add_paragraph()

    # ---------------------------------------------------------------
    #  Title page builder
    # ---------------------------------------------------------------
    def _add_title_page(self, title: str, subtitle: str, description: str):
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.bold      = True
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        sub_para = self.doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub_para.add_run(subtitle)
        sr.bold      = True
        sr.font.size = Pt(28)
        sr.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)
        self.doc.add_paragraph()
        desc_para = self.doc.add_paragraph()
        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc_run = desc_para.add_run(description)
        desc_run.font.size      = Pt(11)
        desc_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        ver_para = self.doc.add_paragraph()
        ver_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = ver_para.add_run(f"Version {PKG_VERSION}")
        vr.font.size      = Pt(12)
        vr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        self.doc.add_page_break()

    # ---------------------------------------------------------------
    #  Core generation (shared between book & appendix)
    # ---------------------------------------------------------------
    def _generate_volume(self, md_files: list, output_path: Path,
                         title: str, subtitle: str, description: str):
        self._reset_doc()
        n = len(md_files)
        print(f"\nGenerating DOCX ({output_path.name}) from {n} content files ...")

        # ---- Title page ----
        self._add_title_page(title, subtitle, description)

        # ---- Table of Contents ----
        _add_toc(self.doc, md_files)

        # ---- Articles ----
        for i, p in enumerate(md_files):
            print(f"  [{i + 1:3d}/{n}] {p.name}")
            self.render_article(p)

        print(f"Writing -> {output_path}")
        self.doc.save(str(output_path))
        kb = output_path.stat().st_size / 1024
        print(f"Done.  {kb:.0f} KB")

    # ---------------------------------------------------------------
    #  Public entry points
    # ---------------------------------------------------------------
    def generate_book(self):
        EXPORT_DIR.mkdir(exist_ok=True)
        all_files = sorted(CONTENT_DIR.rglob("*.md"))
        book_files, _ = _partition_files(all_files)
        if not book_files:
            sys.exit("No book markdown files in content/")
        self._generate_volume(
            book_files,
            EXPORT_DIR / "timeline-book.docx",
            "Paradigm Threat:",
            "The Third Story",
            "A cross-chronology investigation challenging the official timeline "
            "of Earth history \u2014 Scaligerian, Fomenko\u2019s New Chronology, "
            "Saturnian Cosmology, and indigenous traditions.",
        )

    def generate_appendix(self):
        EXPORT_DIR.mkdir(exist_ok=True)
        all_files = sorted(CONTENT_DIR.rglob("*.md"))
        _, appendix_files = _partition_files(all_files)
        if not appendix_files:
            print("No appendix markdown files found — skipping appendix DOCX.")
            return
        self._generate_volume(
            appendix_files,
            EXPORT_DIR / "timeline-appendix.docx",
            "Paradigm Threat:",
            "Appendix & References",
            "Supplementary material, author profiles, investigation cross-references, "
            "and credits for the Paradigm Threat timeline project.",
        )

    def generate(self):
        """Generate both book and appendix DOCX files."""
        self.generate_book()
        self.generate_appendix()


# ---------------------------------------------------------------------------
def main():
    renderer = DOCXRenderer()
    renderer.generate()


if __name__ == "__main__":
    main()
